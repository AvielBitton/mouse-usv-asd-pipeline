#!/usr/bin/env python3
"""Assemble Thesis.docx from sections/*.md + figures + manifest (custom MD->docx renderer)."""
import json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
SEC = REPO / "thesis" / "sections"
MANIFEST = json.loads((REPO / "thesis" / "figure_manifest.json").read_text())
OUT = REPO / "thesis" / "Thesis.docx"

# document order (title page handled programmatically; TOC/LoF inserted after front matter)
ORDER_FRONT = ["01_abstract_he", "02_abstract_en", "03_abbreviations"]
ORDER_BODY = ["10_introduction", "20_background", "30_data", "40_methods", "50_results",
              "60_discussion", "70_reproducibility", "80_limitations", "90_conclusions",
              "95_references", "99_appendix"]


# ---------- figure numbering (pass 1) ----------
def build_fig_numbers():
    order = []
    for stem in ORDER_FRONT + ORDER_BODY:
        txt = (SEC / f"{stem}.md").read_text()
        for m in re.finditer(r"\[\[FIG:([A-Za-z0-9]+)\]\]", txt):
            k = m.group(1)
            if k not in order:
                order.append(k)
    return {k: i + 1 for i, k in enumerate(order)}


FIGNUM = build_fig_numbers()


# ---------- low-level helpers ----------
def set_rtl(par):
    pPr = par._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi"); pPr.append(bidi)
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_runs(par, text):
    """Inline: **bold**, *italic*, `code`, [[REF:Fx]] -> 'Figure N'. """
    text = re.sub(r"\[\[REF:([A-Za-z0-9]+)\]\]",
                  lambda m: f"Figure {FIGNUM.get(m.group(1), '?')}", text)
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = par.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            par.add_run(part)


def add_field(par, instr):
    """Insert a Word field (e.g. TOC) that updates on open."""
    run = par.add_run()._r
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); run.append(fb)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; run.append(it)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); run.append(fs)
    t = OxmlElement("w:t"); t.text = "Update this field to generate the table of contents."; run.append(t)
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); run.append(fe)


def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]; cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in body:
        cells = t.add_row().cells
        for j, val in enumerate(row[:len(header)]):
            p = cells[j].paragraphs[0]; p.text = ""
            add_runs(p, val)
            for r in p.runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph()


def add_figure(doc, fkey):
    info = MANIFEST.get(fkey)
    num = FIGNUM.get(fkey, "?")
    if not info:
        doc.add_paragraph(f"[missing figure {fkey}]")
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(REPO / info["png"]), width=Inches(6.2))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else cap.style
    r = cap.add_run(f"Figure {num}. "); r.bold = True; r.font.size = Pt(9)
    r2 = cap.add_run(info["caption"]); r2.italic = True; r2.font.size = Pt(9)
    doc.add_paragraph()


# ---------- markdown block renderer ----------
def render_md(doc, stem, rtl=False):
    lines = (SEC / f"{stem}.md").read_text().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        # figure directive
        mfig = re.fullmatch(r"\[\[FIG:([A-Za-z0-9]+)\]\]", line.strip())
        if mfig:
            add_figure(doc, mfig.group(1)); i += 1; continue
        # headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3); i += 1; continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2); i += 1; continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            if rtl: set_rtl(h)
            i += 1; continue
        # table block
        if line.lstrip().startswith("|"):
            blk = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                blk.append(lines[i]); i += 1
            if len(blk) >= 2:
                add_table(doc, blk)
            continue
        # bullet list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            continue
        # numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                p = doc.add_paragraph(style="List Number")
                add_runs(p, re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            continue
        # paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, line)
        if rtl: set_rtl(p)
        i += 1


# ---------- title page + front matter ----------
def title_page(doc):
    def center(txt, size, bold=False, italic=False, space=6, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(space)
        return p
    for _ in range(2):
        doc.add_paragraph()
    center("Holon Institute of Technology", 15, bold=True)
    center("Faculty of Sciences — Department of Data Sciences", 12)
    center("M.Sc. Final Project", 12, italic=True, space=28)
    center("Analysis of Ultrasonic Vocalization of Mice for Autism Detection", 20, bold=True, space=10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p); r = p.add_run("ניתוח קולות על-שמע של עכברים לאבחון אוטיזם"); r.font.size = Pt(16); r.bold = True
    p.paragraph_format.space_after = Pt(36)
    center("Chen Aharon   ·   Aviel Bitton", 13, bold=True, space=24)
    center("Advisor: Dr. Dror Lederman", 12.5, bold=True, space=4)
    center("(Holon Institute of Technology, Faculty of Engineering)", 10.5, italic=True, space=14)
    center("In collaboration with Dr. Hava M. Golan", 12)
    center("Biology lead researcher & data owner — Ben-Gurion University of the Negev", 10.5, italic=True, space=36)
    center("2026", 12, bold=True)
    doc.add_page_break()


def toc_section(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    add_field(p, r' TOC \o "1-3" \h \z \u ')
    doc.add_paragraph("(Open in Word and choose “Update Field” to populate the contents and page numbers.)").italic = True
    doc.add_page_break()
    # static list of figures from manifest + numbering
    doc.add_heading("List of Figures", level=1)
    for fkey, num in sorted(FIGNUM.items(), key=lambda kv: kv[1]):
        info = MANIFEST.get(fkey, {})
        para = doc.add_paragraph()
        r = para.add_run(f"Figure {num}. "); r.bold = True; r.font.size = Pt(10)
        para.add_run(info.get("title", fkey)).font.size = Pt(10)
    doc.add_page_break()


# ---------- base styling ----------
def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
    for lvl, sz, col in [("Heading 1", 16, RGBColor(0x0B, 0x3D, 0x66)),
                         ("Heading 2", 13, RGBColor(0x0B, 0x3D, 0x66)),
                         ("Heading 3", 11.5, RGBColor(0x33, 0x33, 0x33))]:
        st = doc.styles[lvl]; st.font.size = Pt(sz); st.font.bold = True; st.font.color.rgb = col
        st.font.name = "Calibri"
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.0)
        sec.left_margin = sec.right_margin = Inches(1.0)


def main():
    doc = Document()
    style_document(doc)
    title_page(doc)
    render_md(doc, "01_abstract_he", rtl=True)
    doc.add_page_break()
    render_md(doc, "02_abstract_en")
    doc.add_page_break()
    render_md(doc, "03_abbreviations")
    doc.add_page_break()
    toc_section(doc)
    for k, stem in enumerate(ORDER_BODY):
        render_md(doc, stem)
        if stem != ORDER_BODY[-1]:
            doc.add_page_break()
    # traceability appendix if present
    trace = REPO / "thesis" / "traceability.md"
    if trace.exists():
        doc.add_page_break()
        # render traceability.md inline
        tmp = SEC / "_traceability.md"
        tmp.write_text(trace.read_text())
        render_md(doc, "_traceability")
        tmp.unlink()
    doc.save(OUT)
    # report
    figs = len([1 for s in ORDER_FRONT + ORDER_BODY for _ in re.finditer(r"\[\[FIG:", (SEC / f"{s}.md").read_text())])
    print(f"Saved {OUT.relative_to(REPO)}  | figures placed: {figs} | figure numbers: {len(FIGNUM)}")


if __name__ == "__main__":
    main()
